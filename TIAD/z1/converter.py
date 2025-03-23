import openpyxl
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert


def read_excel(file_path):
    excel = openpyxl.load_workbook(file_path)
    table = excel.active
    data = []
    for row in table.iter_rows(values_only=True):
        data.append(row)
    return data


def page_numbering(doc):
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._r.append(instrText)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)


def save_to_docx_lines(data, out, title, space, align, header_b, header_i, value_b, value_i, page_num):
    doc = Document()

    doc.add_heading(title, level=1)
    doc.add_paragraph("\n")

    para_align_dict = {"Wyrównaj do lewej": WD_PARAGRAPH_ALIGNMENT.LEFT, "Wyśrodkowany": WD_PARAGRAPH_ALIGNMENT.CENTER,
                 "Wyrównaj do prawej": WD_PARAGRAPH_ALIGNMENT.RIGHT}


    headers = data[0]

    for row in data[1:]:
        for i, (col_name, cell_val) in enumerate(zip(headers, row)):
            para = doc.add_paragraph()
            run = para.add_run(f"{col_name}: ") # w python-docx run to fragment tekstu w paragrafie
            run.bold = header_b
            run.italic = header_i
            run.font.size = Pt(12)
            run = para.add_run(f"{cell_val}")
            run.bold = value_b
            run.italic = value_i
            run.font.size = Pt(12)
            para.alignment = para_align_dict.get(align, WD_PARAGRAPH_ALIGNMENT.LEFT)

        # przerwa pomiedzy linijkami z excela
        doc.add_paragraph().paragraph_format.space_after = Pt(space)

    if page_num:
        page_numbering(doc)

    doc.save(out)


def save_to_docx_table(data, out, title, align, header_b, header_i, value_b, value_i, page_num):
    doc = Document()

    #tytul
    doc.add_heading(title, level=1)
    doc.add_paragraph("\n")

    para_align_dict = {"Wyrównaj do lewej": WD_PARAGRAPH_ALIGNMENT.LEFT, "Wyśrodkowany": WD_PARAGRAPH_ALIGNMENT.CENTER,
                 "Wyrównaj do prawej": WD_PARAGRAPH_ALIGNMENT.RIGHT}

    table = doc.add_table(rows=1, cols=len(data[0]))
    table.style = "Table Grid"

    # naglowki
    headers = table.rows[0].cells
    for i, header in enumerate(data[0]):
        headers[i].text = header
        headers[i].paragraphs[0].alignment = para_align_dict.get(align, WD_PARAGRAPH_ALIGNMENT.LEFT)
        if header_b:
            headers[i].paragraphs[0].runs[0].bold = True
        if header_i:
            headers[i].paragraphs[0].runs[0].italic = True

    # dane
    for row in data[1:]:
        row_values = table.add_row().cells
        for i, cell_val in enumerate(row):
            row_values[i].text = str(cell_val)
            row_values[i].paragraphs[0].alignment = para_align_dict.get(align, WD_PARAGRAPH_ALIGNMENT.LEFT)
            if value_b:
                row_values[i].paragraphs[0].runs[0].bold = True
            if value_i:
                row_values[i].paragraphs[0].runs[0].italic = True

    if page_num:
        page_numbering(doc)

    doc.save(out)


def convert_excel(file_path, docx_out, pdf_out, title, space, align, header_b, header_i, value_b, value_i, page_num, format):
    data = read_excel(file_path)

    if format == "lines":
        save_to_docx_lines(data, docx_out, title, space, align, header_b, header_i,
                                  value_b, value_i, page_num)
    else:
        save_to_docx_table(data, docx_out, title, align, header_b, header_i, value_b,
                           value_i, page_num)

    # to wyświetla pasek postępu w terminalu
    convert(docx_out, pdf_out)